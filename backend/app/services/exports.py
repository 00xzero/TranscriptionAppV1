from io import BytesIO
from typing import List, Optional
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer


def format_duration(seconds: int) -> str:
    """
    Convert seconds to human-readable duration format.
    
    Examples:
        3898 -> "1h 4m 58s"
        125 -> "2m 5s"
        45 -> "45s"
    """
    if seconds < 0:
        seconds = 0
    
    hours = seconds // 3600
    minutes = (seconds % 3600) // 60
    secs = seconds % 60
    
    parts = []
    if hours > 0:
        parts.append(f"{hours}h")
    if minutes > 0:
        parts.append(f"{minutes}m")
    if secs > 0 or not parts:  # Always show seconds if it's the only unit
        parts.append(f"{secs}s")
    
    return " ".join(parts)


def generate_docx(
    project_title: str,
    chunks: List[dict],
    speakers_map: dict,
    transcription_date: datetime,
    duration_seconds: Optional[int] = None,
) -> BytesIO:
    """
    Generate a DOCX file from transcript chunks following PRD format.
    
    Args:
        project_title: Title of the project/transcript
        chunks: List of transcript chunks with speaker_id, start_ms, end_ms, text
        speakers_map: Dict mapping speaker_id to speaker info (label, color)
        transcription_date: Date when transcription was created
        duration_seconds: Optional duration of the audio in seconds
    
    Returns:
        BytesIO buffer containing the DOCX file
    """
    doc = Document()
    
    # Title - centered, 16pt bold
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run(project_title or "Transcript")
    title_run.bold = True
    title_run.font.size = Pt(16)
    
    # Metadata block - centered, smaller font, gray
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Date of Transcription
    date_str = transcription_date.strftime("%B %d, %Y, %I:%M%p")
    meta_run = meta_para.add_run(date_str)
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Duration (if available)
    if duration_seconds is not None:
        meta_para.add_run("\n")
        duration_run = meta_para.add_run(format_duration(duration_seconds))
        duration_run.font.size = Pt(10)
        duration_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Spacer
    doc.add_paragraph()
    
    # Transcript body - group chunks by speaker turns
    current_speaker_id = None
    current_paragraph = None
    
    for chunk in chunks:
        speaker_id = chunk.get("speaker_id")
        speaker_label = speakers_map.get(speaker_id, {}).get("label", "Unknown Speaker")
        start_ms = chunk.get("start_ms", 0)
        end_ms = chunk.get("end_ms", 0)
        text = chunk.get("text", "")
        
        # Check if we need a new paragraph (new speaker)
        if speaker_id != current_speaker_id:
            current_speaker_id = speaker_id
            current_paragraph = doc.add_paragraph()
            
            # Speaker label - bold
            speaker_run = current_paragraph.add_run(f"{speaker_label}\n")
            speaker_run.bold = True
            speaker_run.font.size = Pt(12)
        
        # Add timestamp and text on same line
        if current_paragraph:
            # Timestamp in gray
            timestamp_str = f"{ms_to_timestamp(start_ms)}\n"
            ts_run = current_paragraph.add_run(timestamp_str)
            ts_run.font.size = Pt(10)
            ts_run.font.color.rgb = RGBColor(100, 100, 100)
            
            # Spoken text
            text_run = current_paragraph.add_run(text)
            text_run.font.size = Pt(11)
            
            # Add spacing between chunks
            current_paragraph.add_run("\n")
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_vtt(
    chunks: List[dict],
    speakers_map: dict,
    project_id: str = "transcript"
) -> str:
    """
    Generate a WebVTT file from transcript chunks following PRD format.
    
    Args:
        chunks: List of transcript chunks
        speakers_map: Dict mapping speaker_id to speaker info
        project_id: Project ID for cue identifiers
    
    Returns:
        VTT content as string
    """
    lines = ["WEBVTT", ""]
    
    for idx, chunk in enumerate(chunks):
        speaker_id = chunk.get("speaker_id")
        speaker_label = speakers_map.get(speaker_id, {}).get("label", "Speaker")
        
        start_ms = chunk.get("start_ms", 0)
        end_ms = chunk.get("end_ms", 0)
        text = chunk.get("text", "")
        
        # Cue identifier format: {project_id}/{index}
        cue_id = f"{project_id}/{idx}"
        
        # VTT timestamps
        start_vtt = ms_to_vtt_timestamp(start_ms)
        end_vtt = ms_to_vtt_timestamp(end_ms)
        
        # Build cue
        lines.append(cue_id)
        lines.append(f"{start_vtt} --> {end_vtt}")
        lines.append(f"<v {speaker_label}>{text}</v>")
        lines.append("")
    
    return "\n".join(lines)


def generate_pdf(
    project_title: str,
    chunks: List[dict],
    speakers_map: dict,
    transcription_date: datetime,
    duration_seconds: Optional[int] = None,
) -> BytesIO:
    """
    Generate a PDF file matching DOCX structure following PRD format.
    
    Args:
        project_title: Title of the project/transcript
        chunks: List of transcript chunks
        speakers_map: Dict mapping speaker_id to speaker info
        transcription_date: Date when transcription was created
        duration_seconds: Optional duration of the audio in seconds
    
    Returns:
        BytesIO buffer containing the PDF file
    """
    buffer = BytesIO()
    
    # Create PDF with 1-inch margins
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=1*inch,
        rightMargin=1*inch,
        topMargin=1*inch,
        bottomMargin=1*inch,
    )
    
    # Container for PDF elements
    story = []
    
    # Define styles
    styles = getSampleStyleSheet()
    
    # Title style - centered, bold, larger
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        textColor='black',
        spaceAfter=6,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold',
    )
    
    # Metadata style - centered, smaller, gray
    meta_style = ParagraphStyle(
        'CustomMeta',
        parent=styles['Normal'],
        fontSize=10,
        textColor='#808080',
        alignment=TA_CENTER,
        spaceAfter=12,
    )
    
    # Speaker label style - bold
    speaker_style = ParagraphStyle(
        'CustomSpeaker',
        parent=styles['Normal'],
        fontSize=12,
        fontName='Helvetica-Bold',
        spaceAfter=2,
    )
    
    # Timestamp style - smaller, gray
    timestamp_style = ParagraphStyle(
        'CustomTimestamp',
        parent=styles['Normal'],
        fontSize=10,
        textColor='#646464',
        spaceAfter=2,
    )
    
    # Body text style
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=8,
    )
    
    # Add title
    story.append(Paragraph(project_title or "Transcript", title_style))
    
    # Add metadata
    date_str = transcription_date.strftime("%B %d, %Y, %I:%M%p")
    meta_text = date_str
    if duration_seconds is not None:
        meta_text += f"<br/>{format_duration(duration_seconds)}"
    story.append(Paragraph(meta_text, meta_style))
    
    # Add spacer
    story.append(Spacer(1, 0.2*inch))
    
    # Add transcript chunks
    current_speaker_id = None
    
    for chunk in chunks:
        speaker_id = chunk.get("speaker_id")
        speaker_label = speakers_map.get(speaker_id, {}).get("label", "Unknown Speaker")
        start_ms = chunk.get("start_ms", 0)
        text = chunk.get("text", "")
        
        # New speaker - add speaker label
        if speaker_id != current_speaker_id:
            current_speaker_id = speaker_id
            story.append(Paragraph(speaker_label, speaker_style))
        
        # Add timestamp
        timestamp_str = ms_to_timestamp(start_ms)
        story.append(Paragraph(timestamp_str, timestamp_style))
        
        # Add text
        # Escape XML special characters for ReportLab
        safe_text = (text
                     .replace("&", "&amp;")
                     .replace("<", "&lt;")
                     .replace(">", "&gt;"))
        story.append(Paragraph(safe_text, body_style))
    
    # Build PDF
    doc.build(story)
    
    buffer.seek(0)
    return buffer


def ms_to_timestamp(ms: int) -> str:
    """
    Convert milliseconds to MM:SS or H:MM:SS format.
    
    Examples:
        4205 -> "0:04"
        65000 -> "1:05"
        3665000 -> "1:01:05"
    """
    total_sec = ms // 1000
    s = total_sec % 60
    m = (total_sec // 60) % 60
    h = total_sec // 3600
    
    if h > 0:
        return f"{h}:{m:02d}:{s:02d}"
    return f"{m}:{s:02d}"


def ms_to_vtt_timestamp(ms: int) -> str:
    """
    Convert milliseconds to VTT timestamp format (HH:MM:SS.mmm).
    
    Examples:
        4205 -> "00:00:04.205"
        65000 -> "00:01:05.000"
    """
    total_sec = ms // 1000
    millis = ms % 1000
    s = total_sec % 60
    m = (total_sec // 60) % 60
    h = total_sec // 3600
    
    return f"{h:02d}:{m:02d}:{s:02d}.{millis:03d}"
